import base64
import os
import tempfile
import unittest
from pathlib import Path
from unittest import mock

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase
from django.urls import reverse

from sc_admin_portal.news_storage import NewsStorageError
from sc_admin_portal.tests.test_news_docx_import import _PIL_AVAILABLE

ADMIN_EMAIL = "admin@example.com"


class AdminPortalNewsPublishTests(TestCase):
    def setUp(self):
        self.user_model = get_user_model()
        self.authorized_user = self.user_model.objects.create_user(
            username="authorized",
            email=ADMIN_EMAIL,
            password="pass",
        )

    @mock.patch.dict(os.environ, {"SC_ADMIN_EMAIL": ADMIN_EMAIL})
    @mock.patch("sc_admin_portal.views.create_news_post")
    @mock.patch("sc_admin_portal.views.oracle_proc.list_recent_jobs", return_value=[])
    @mock.patch("sc_admin_portal.views.oracle_proc.list_pending_approvals", return_value=[])
    @mock.patch("sc_admin_portal.views.oracle_proc.list_recent_decisions", return_value=[])
    @mock.patch("sc_admin_portal.views.oracle_proc.list_recent_research_requests", return_value=[])
    def test_publish_news_creates_post(
        self,
        research_mock,
        decisions_mock,
        approvals_mock,
        jobs_mock,
        create_news_post,
    ):
        create_news_post.return_value = {"id": "NEWS_ITEMS:12", "title": "Hello"}
        self.client.force_login(self.authorized_user)
        response = self.client.post(
            reverse("sc_admin_portal:dashboard"),
            {
                "action": "publish_news",
                "headline": "Hello",
                "tags": "AI, Governance",
                "body_html": "<p>Body</p>",
                "confirm_no_images": "on",
            },
        )
        self.assertEqual(response.status_code, 302)
        self.assertIn("/news/NEWS_ITEMS:12/", response["Location"])
        create_news_post.assert_called_once()

    @mock.patch.dict(os.environ, {"SC_ADMIN_EMAIL": ADMIN_EMAIL})
    @mock.patch("sc_admin_portal.views.create_news_asset", return_value=55)
    def test_news_asset_upload_returns_location(self, create_news_asset):
        self.client.force_login(self.authorized_user)
        upload = SimpleUploadedFile(
            "chart.png", b"fake-image", content_type="image/png"
        )
        response = self.client.post(
            reverse("sc_admin_portal:news_asset_upload"),
            {"file": upload},
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["asset_id"], 55)
        self.assertIn("http://testserver/news/assets/55/", payload["location"])

    @mock.patch.dict(os.environ, {"SC_ADMIN_EMAIL": ADMIN_EMAIL})
    @mock.patch(
        "sc_admin_portal.views.create_news_asset",
        side_effect=NewsStorageError(
            "NEWS_ASSETS table is missing. Apply migration V0004__news_rich_body.sql.",
            code="missing_news_assets",
        ),
    )
    def test_news_asset_upload_surfaces_missing_table(self, create_news_asset):
        self.client.force_login(self.authorized_user)
        upload = SimpleUploadedFile(
            "chart.png", b"fake-image", content_type="image/png"
        )
        response = self.client.post(
            reverse("sc_admin_portal:news_asset_upload"),
            {"file": upload},
        )
        self.assertEqual(response.status_code, 503)
        payload = response.json()
        self.assertEqual(payload["error"], "missing_news_assets")
        self.assertIn("V0004__news_rich_body.sql", payload["message"])

    @mock.patch.dict(os.environ, {"SC_ADMIN_EMAIL": ADMIN_EMAIL})
    @mock.patch(
        "sc_admin_portal.views.create_news_post",
        side_effect=NewsStorageError(
            "BODY_HTML column is missing. Apply migration V0004__news_rich_body.sql.",
            code="missing_body_html",
        ),
    )
    @mock.patch("sc_admin_portal.views.oracle_proc.list_recent_jobs", return_value=[])
    @mock.patch("sc_admin_portal.views.oracle_proc.list_pending_approvals", return_value=[])
    @mock.patch("sc_admin_portal.views.oracle_proc.list_recent_decisions", return_value=[])
    @mock.patch("sc_admin_portal.views.oracle_proc.list_recent_research_requests", return_value=[])
    def test_publish_news_surfaces_missing_body_html(
        self,
        research_mock,
        decisions_mock,
        approvals_mock,
        jobs_mock,
        create_news_post,
    ):
        self.client.force_login(self.authorized_user)
        response = self.client.post(
            reverse("sc_admin_portal:dashboard"),
            {
                "action": "publish_news",
                "headline": "Hello",
                "tags": "AI, Governance",
                "body_html": "<p>Body</p>",
                "confirm_no_images": "on",
            },
        )
        self.assertEqual(response.status_code, 200)
        content = response.content.decode("utf-8")
        self.assertIn("BODY_HTML column is missing", content)

    @mock.patch.dict(os.environ, {"SC_ADMIN_EMAIL": ADMIN_EMAIL})
    @mock.patch("sc_admin_portal.views.create_news_post")
    @mock.patch("sc_admin_portal.views.oracle_proc.list_recent_jobs", return_value=[])
    @mock.patch("sc_admin_portal.views.oracle_proc.list_pending_approvals", return_value=[])
    @mock.patch("sc_admin_portal.views.oracle_proc.list_recent_decisions", return_value=[])
    @mock.patch("sc_admin_portal.views.oracle_proc.list_recent_research_requests", return_value=[])
    def test_publish_news_requires_no_image_confirm(
        self,
        research_mock,
        decisions_mock,
        approvals_mock,
        jobs_mock,
        create_news_post,
    ):
        self.client.force_login(self.authorized_user)
        response = self.client.post(
            reverse("sc_admin_portal:dashboard"),
            {
                "action": "publish_news",
                "headline": "Hello",
                "tags": "AI, Governance",
                "body_html": "<p>Body</p>",
            },
        )
        self.assertEqual(response.status_code, 200)
        content = response.content.decode("utf-8")
        self.assertIn("Confirm that no images are included", content)
        create_news_post.assert_not_called()

    @unittest.skipUnless(_PIL_AVAILABLE, "Pillow not installed; skipping image upload test.")
    @mock.patch.dict(os.environ, {"SC_ADMIN_EMAIL": ADMIN_EMAIL})
    @mock.patch("sc_admin_portal.views.create_news_asset", return_value=55)
    @mock.patch("sc_admin_portal.views.create_news_post")
    @mock.patch("sc_admin_portal.views.oracle_proc.list_recent_jobs", return_value=[])
    @mock.patch("sc_admin_portal.views.oracle_proc.list_pending_approvals", return_value=[])
    @mock.patch("sc_admin_portal.views.oracle_proc.list_recent_decisions", return_value=[])
    @mock.patch("sc_admin_portal.views.oracle_proc.list_recent_research_requests", return_value=[])
    def test_publish_news_with_docx_imports_images(
        self,
        research_mock,
        decisions_mock,
        approvals_mock,
        jobs_mock,
        create_news_post,
        create_news_asset,
    ):
        from docx import Document

        create_news_post.return_value = {"id": "NEWS_ITEMS:55", "title": "Docx"}
        self.client.force_login(self.authorized_user)

        with tempfile.TemporaryDirectory() as tmpdir:
            image_path = Path(tmpdir) / "tiny.png"
            image_bytes = base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMB/6n8lHcAAAAASUVORK5CYII="
            )
            image_path.write_bytes(image_bytes)

            doc = Document()
            doc.add_heading("Docx headline", level=1)
            doc.add_picture(str(image_path))
            docx_path = Path(tmpdir) / "news.docx"
            doc.save(docx_path)

            with open(docx_path, "rb") as handle:
                upload = SimpleUploadedFile(
                    "news.docx",
                    handle.read(),
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

            response = self.client.post(
                reverse("sc_admin_portal:dashboard"),
                {
                    "action": "publish_news",
                    "headline": "Override headline",
                    "tags": "AI",
                },
                files={"docx_file": upload},
            )

        self.assertEqual(response.status_code, 302)
        args, kwargs = create_news_post.call_args
        self.assertIn("/news/assets/55/", kwargs["body_html"])
        create_news_asset.assert_called()
