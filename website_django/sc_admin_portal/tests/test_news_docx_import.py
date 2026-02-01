import base64
import tempfile
from pathlib import Path
import unittest

from django.test import SimpleTestCase
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from sc_admin_portal.docx_import import build_news_body_from_docx

try:
    from PIL import Image  # noqa: F401

    _PIL_AVAILABLE = True
except ImportError:  # pragma: no cover - optional dependency for image tests
    _PIL_AVAILABLE = False


class DocxImportTests(SimpleTestCase):
    def _add_hyperlink(self, paragraph, url: str, text_parts: list[str]) -> None:
        rel_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), rel_id)
        for text in text_parts:
            run = OxmlElement("w:r")
            text_el = OxmlElement("w:t")
            text_el.text = text
            run.append(text_el)
            hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def test_docx_import_extracts_headline_and_body(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            doc = Document()
            doc.add_heading("DOCX headline", level=1)
            doc.add_paragraph("First paragraph.")
            path = Path(tmpdir) / "news.docx"
            doc.save(path)

            headline, body_html = build_news_body_from_docx(
                str(path),
                asset_uploader=lambda *args, **kwargs: 1,
            )

            self.assertEqual(headline, "DOCX headline")
            self.assertIn("<p>First paragraph.</p>", body_html)

    def test_docx_import_renders_tables(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            doc = Document()
            doc.add_heading("Table headline", level=1)
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Cell A"
            table.cell(0, 1).text = "Cell B"
            path = Path(tmpdir) / "news.docx"
            doc.save(path)

            headline, body_html = build_news_body_from_docx(
                str(path),
                asset_uploader=lambda *args, **kwargs: 1,
            )

            self.assertEqual(headline, "Table headline")
            self.assertIn("<table>", body_html)
            self.assertIn("<td>Cell A</td>", body_html)

    @unittest.skipUnless(_PIL_AVAILABLE, "Pillow not installed; skipping image import test.")
    def test_docx_import_inserts_images(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            image_path = Path(tmpdir) / "tiny.png"
            image_bytes = base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMB/6n8lHcAAAAASUVORK5CYII="
            )
            image_path.write_bytes(image_bytes)

            doc = Document()
            doc.add_heading("Image headline", level=1)
            doc.add_picture(str(image_path))
            path = Path(tmpdir) / "news.docx"
            doc.save(path)

            uploads = []
            stats: dict[str, int] = {}

            def uploader(file_name, mime_type, file_bytes):
                uploads.append((file_name, mime_type, file_bytes))
                return 42

            headline, body_html = build_news_body_from_docx(
                str(path),
                asset_uploader=uploader,
                stats=stats,
            )

            self.assertEqual(headline, "Image headline")
            self.assertIn("/news/assets/42/", body_html)
            self.assertTrue(uploads)
            self.assertEqual(stats.get("images_found"), 1)
            self.assertEqual(stats.get("images_uploaded"), 1)

    def test_docx_import_preserves_hyperlinks(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            doc = Document()
            doc.add_heading("Hyperlink headline", level=1)
            paragraph = doc.add_paragraph()
            self._add_hyperlink(
                paragraph,
                "/tech100/company/MSFT/",
                ["Microsoft ", "(MSFT)"],
            )
            paragraph.add_run(": AI governance summary.")

            paragraph_two = doc.add_paragraph()
            paragraph_two.add_run("Alphabet (Google) ")
            self._add_hyperlink(
                paragraph_two,
                "/tech100/company/GOOGL/",
                ["(GOOGL)"],
            )

            path = Path(tmpdir) / "news.docx"
            doc.save(path)

            headline, body_html = build_news_body_from_docx(
                str(path),
                asset_uploader=lambda *args, **kwargs: 1,
            )

            self.assertEqual(headline, "Hyperlink headline")
            self.assertIn('href="/tech100/company/MSFT/"', body_html)
            self.assertIn("Microsoft", body_html)
            self.assertIn("MSFT", body_html)
            self.assertIn('href="/tech100/company/GOOGL/"', body_html)
