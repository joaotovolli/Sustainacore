"""DOCX and chart generation."""
from __future__ import annotations

import logging
import os
from typing import Any, Dict, List

LOGGER = logging.getLogger("research_generator.docx")


def _ensure_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)


def _render_chart(chart_data: Dict[str, Any], output_dir: str, report_key: str) -> str:
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception as exc:
        raise RuntimeError("matplotlib_missing") from exc

    _ensure_dir(output_dir)
    chart_path = os.path.join(output_dir, f"{report_key}_chart.png")
    chart_type = chart_data.get("type", "line")
    x_vals = chart_data.get("x") or []
    series = chart_data.get("series") or []
    y_vals = chart_data.get("y") or []

    plt.figure(figsize=(6.8, 3.6))
    if chart_type == "bar":
        if series:
            width = 0.8 / max(len(series), 1)
            indices = list(range(len(x_vals)))
            for idx, entry in enumerate(series):
                values = entry.get("values") or []
                offsets = [i + idx * width for i in indices]
                plt.bar(offsets, values, width=width, label=entry.get("name") or f"Series {idx+1}")
            plt.xticks(
                [i + width * (len(series) - 1) / 2 for i in indices],
                x_vals,
                rotation=45,
                ha="right",
            )
            if len(series) > 1:
                plt.legend(fontsize="small")
        else:
            plt.bar(x_vals, y_vals)
            plt.xticks(rotation=45, ha="right")
    elif chart_type == "box":
        data = [entry.get("values") or [] for entry in series]
        labels = [entry.get("name") or "" for entry in series]
        plt.boxplot(data, labels=labels)
    elif chart_type == "hist":
        data = [entry.get("values") or [] for entry in series]
        labels = [entry.get("name") or "" for entry in series]
        if data:
            plt.hist(data, bins=10, label=labels, alpha=0.7)
            if len(labels) > 1:
                plt.legend(fontsize="small")
    else:
        if series:
            for idx, entry in enumerate(series):
                values = entry.get("values") or []
                plt.plot(x_vals, values, marker="o", label=entry.get("name") or f"Series {idx+1}")
            if len(series) > 1:
                plt.legend(fontsize="small")
        else:
            plt.plot(x_vals, y_vals, marker="o")
    plt.title(chart_data.get("title") or "")
    plt.xticks(rotation=45, ha="right")
    plt.tight_layout()
    plt.savefig(chart_path, dpi=150)
    plt.close()
    return chart_path


def _add_table(document, table_rows: List[Dict[str, Any]]) -> None:
    if not table_rows:
        return
    columns = list(table_rows[0].keys())
    table = document.add_table(rows=1, cols=len(columns))
    hdr_cells = table.rows[0].cells
    for idx, col in enumerate(columns):
        hdr_cells[idx].text = str(col)
    for row in table_rows:
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            cells[idx].text = str(row.get(col, ""))


def build_docx(
    draft: Dict[str, Any],
    bundle: Dict[str, Any],
    report_key: str,
    output_dir: str,
) -> Dict[str, Any]:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python_docx_missing") from exc

    chart_path = _render_chart(bundle.get("chart_data", {}), output_dir, report_key)

    document = Document()
    document.add_heading(draft.get("headline") or "Research Update", level=0)
    dek = draft.get("dek")
    if dek:
        para = document.add_paragraph(str(dek))
        if para.runs:
            para.runs[0].italic = True

    charts = bundle.get("docx_charts") or []
    tables = bundle.get("docx_tables") or []
    outline = draft.get("outline") or []
    chart_paths: Dict[int, str] = {}
    for idx, chart in enumerate(charts, start=1):
        chart_paths[idx] = _render_chart(chart, output_dir, f"{report_key}_{idx}")

    if not outline:
        outline = [{"type": "paragraph", "text": p} for p in draft.get("paragraphs", [])]
        for idx in range(1, len(charts) + 1):
            outline.append({"type": "figure", "id": idx})
        for idx in range(1, len(tables) + 1):
            outline.append({"type": "table", "id": idx})

    inserted_figs = set()
    inserted_tables = set()
    for item in outline:
        if item.get("type") == "paragraph":
            document.add_paragraph(item.get("text") or "")
        elif item.get("type") == "figure":
            fig_id = int(item.get("id") or 0)
            if 0 < fig_id <= len(charts):
                caption = charts[fig_id - 1].get("caption") or f"Figure {fig_id}. Chart"
                if not str(caption).lower().startswith("figure"):
                    caption = f"Figure {fig_id}. {caption}"
                document.add_paragraph(caption)
                path = chart_paths.get(fig_id)
                if path and os.path.exists(path):
                    document.add_picture(path)
                inserted_figs.add(fig_id)
        elif item.get("type") == "table":
            tbl_id = int(item.get("id") or 0)
            if 0 < tbl_id <= len(tables):
                title = tables[tbl_id - 1].get("title") or f"Table {tbl_id}. Table"
                if not title.lower().startswith("table"):
                    title = f"Table {tbl_id}. {title}"
                document.add_paragraph(title)
                _add_table(document, tables[tbl_id - 1].get("rows") or [])
                inserted_tables.add(tbl_id)

    for fig_id in range(1, len(charts) + 1):
        if fig_id in inserted_figs:
            continue
        caption = charts[fig_id - 1].get("caption") or f"Figure {fig_id}. Chart"
        if not str(caption).lower().startswith("figure"):
            caption = f"Figure {fig_id}. {caption}"
        document.add_paragraph(caption)
        path = chart_paths.get(fig_id)
        if path and os.path.exists(path):
            document.add_picture(path)

    for tbl_id in range(1, len(tables) + 1):
        if tbl_id in inserted_tables:
            continue
        title = tables[tbl_id - 1].get("title") or f"Table {tbl_id}. Table"
        if not title.lower().startswith("table"):
            title = f"Table {tbl_id}. {title}"
        document.add_paragraph(title)
        _add_table(document, tables[tbl_id - 1].get("rows") or [])

    document.add_paragraph(
        "Disclaimer: Research and education only; not investment advice."
    )
    document.add_paragraph(f"Methodology: {bundle.get('methodology_url')}")

    output_name = f"{report_key}.docx"
    output_path = os.path.join(output_dir, output_name)
    document.save(output_path)

    with open(output_path, "rb") as handle:
        docx_bytes = handle.read()

    first_chart = ""
    if chart_paths:
        first_chart = chart_paths[min(chart_paths.keys())]
    return {
        "docx_bytes": docx_bytes,
        "docx_name": output_name,
        "chart_path": first_chart,
    }
